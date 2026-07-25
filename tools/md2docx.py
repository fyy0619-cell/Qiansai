# -*- coding: utf-8 -*-
"""把答辩说辞 Markdown 转成排版良好的 .docx，再用 Word 导出 .pdf。中文字体安全。"""
import os, re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK = "微软雅黑"
MONO = "Consolas"

def set_run_font(run, name=CJK, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

TOKEN = re.compile(r'(\*\*.+?\*\*|`.+?`)')

def add_inline(p, text, base_size=10.5, base_color=None, base_bold=False):
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_run_font(run, name=MONO, size=base_size-0.5, color=(0xC7, 0x25, 0x4E))
        else:
            run = p.add_run(part)
            set_run_font(run, size=base_size, bold=base_bold, color=base_color)

def main(md_path, docx_path, title):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = CJK
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), CJK)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        raw = line.rstrip()
        stripped = raw.strip()

        # table block
        if stripped.startswith('|') and i + 1 < n and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i+1]):
            header = [c.strip() for c in stripped.strip('|').split('|')]
            rows = []
            i += 2
            while i < n and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Table Grid'
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].clear()
                add_inline(cell.paragraphs[0], h, base_size=10, base_bold=True)
                shade_cell(cell, 'D9E2F3')
            for r in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    txt = r[j] if j < len(r) else ''
                    cells[j].paragraphs[0].clear()
                    add_inline(cells[j].paragraphs[0], txt, base_size=9.5)
            doc.add_paragraph()
            continue

        if raw == '':
            i += 1
            continue

        # horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'BBBBBB')
            pbdr.append(bottom); pPr.append(pbdr)
            i += 1
            continue

        # headings
        m = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1)); text = m.group(2)
            p = doc.add_paragraph()
            p.space_before = Pt(6)
            sizes = {1: 19, 2: 15, 3: 12.5, 4: 11}
            colors = {1: (0x1F, 0x38, 0x64), 2: (0x1F, 0x4E, 0x79), 3: (0x2E, 0x54, 0x96), 4: (0x40, 0x40, 0x40)}
            p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, text, base_size=sizes[level], base_color=colors[level], base_bold=True)
            i += 1
            continue

        # blockquote (可能多行)
        if stripped.startswith('>'):
            qlines = []
            while i < n and lines[i].strip().startswith('>'):
                qlines.append(re.sub(r'^\s*>\s?', '', lines[i]))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18')
            left.set(qn('w:space'), '8'); left.set(qn('w:color'), 'B4C6E7')
            pbdr.append(left); pPr.append(pbdr)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F2F5FB')
            pPr.append(shd)
            add_inline(p, '\n'.join(qlines), base_size=9.5, base_color=(0x44, 0x44, 0x44))
            continue

        # checkbox list
        mcb = re.match(r'^-\s*\[\s*\]\s*(.*)$', stripped)
        if mcb:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run('☐  '); set_run_font(run, size=10.5)
            add_inline(p, mcb.group(1))
            i += 1
            continue

        # bullet list
        mb = re.match(r'^[-*]\s+(.*)$', stripped)
        if mb:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, mb.group(1))
            i += 1
            continue

        # numbered list
        mn = re.match(r'^\d+\.\s+(.*)$', stripped)
        if mn:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, mn.group(1))
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print("DOCX saved:", docx_path)

if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else '')
